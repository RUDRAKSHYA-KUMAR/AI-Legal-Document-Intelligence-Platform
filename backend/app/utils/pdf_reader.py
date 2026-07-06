"""
app/utils/pdf_reader.py

Document reading utilities for the AI Legal Documentation Platform.

Supported Formats:
- PDF
- DOCX
- TXT

This module extracts:
- Full Text
- Metadata
- Page Information
- Preview

Future Compatible:
- RAG
- Gemini Analysis
- Dashboard
- Chat
"""

from pathlib import Path
import logging

import fitz  # PyMuPDF
import docx as python_docx


# ======================================================
# Logger
# ======================================================

logger = logging.getLogger(__name__)


# ======================================================
# Helpers
# ======================================================

def clean_text(text: str) -> str:
    """
    Clean extracted text.

    Removes:
    - Extra spaces
    - Empty lines
    """

    return "\n".join(
        line.strip()
        for line in text.splitlines()
        if line.strip()
    )


def get_text_preview(text: str, max_chars: int = 500) -> str:
    """
    Return a short preview of extracted text.
    """

    cleaned = " ".join(text.split())

    if len(cleaned) <= max_chars:
        return cleaned

    return cleaned[:max_chars] + "..."


# ======================================================
# PDF Reader
# ======================================================

def extract_text_from_pdf(path: str) -> dict:

    file_path = Path(path)

    if not file_path.exists():
        raise FileNotFoundError(f"PDF not found: {path}")

    logger.info(f"Reading PDF: {path}")

    try:

        doc = fitz.open(file_path)

        pages = []

        for page_number, page in enumerate(doc, start=1):

            text = clean_text(
                page.get_text("text")
            )

            pages.append({
                "page_number": page_number,
                "text": text,
                "char_count": len(text)
            })

        full_text = "\n\n".join(
            f"[Page {page['page_number']}]\n{page['text']}"
            for page in pages
            if page["text"]
        )

        metadata = doc.metadata or {}

        doc.close()

        return {

            "text": full_text,

            "preview": get_text_preview(full_text),

            "page_count": len(pages),

            "pages": pages,

            "metadata": {

                "title": metadata.get("title", ""),

                "author": metadata.get("author", ""),

                "creator": metadata.get("creator", ""),

                "subject": metadata.get("subject", "")
            }

        }

    except Exception as e:

        logger.exception("PDF Extraction Failed")

        raise RuntimeError(
            f"Failed to read PDF: {e}"
        )


# ======================================================
# DOCX Reader
# ======================================================

def extract_text_from_docx(path: str) -> dict:

    file_path = Path(path)

    if not file_path.exists():
        raise FileNotFoundError(f"DOCX not found: {path}")

    logger.info(f"Reading DOCX: {path}")

    try:

        document = python_docx.Document(file_path)

        paragraphs = [

            clean_text(p.text)

            for p in document.paragraphs

            if p.text.strip()

        ]

        table_rows = []

        for table in document.tables:

            for row in table.rows:

                row_text = " | ".join(

                    cell.text.strip()

                    for cell in row.cells

                    if cell.text.strip()

                )

                if row_text:
                    table_rows.append(row_text)

        all_text = paragraphs

        if table_rows:

            all_text.append("[Tables]")

            all_text.extend(table_rows)

        full_text = "\n\n".join(all_text)

        props = document.core_properties

        word_count = len(full_text.split())

        return {

            "text": full_text,

            "preview": get_text_preview(full_text),

            "page_count": max(1, word_count // 500),

            "pages": [],

            "metadata": {

                "title": props.title or "",

                "author": props.author or "",

                "subject": props.subject or "",

                "created": str(props.created)
                if props.created else ""

            }

        }

    except Exception as e:

        logger.exception("DOCX Extraction Failed")

        raise RuntimeError(
            f"Failed to read DOCX: {e}"
        )


# ======================================================
# TXT Reader
# ======================================================

def extract_text_from_txt(path: str) -> dict:

    file_path = Path(path)

    if not file_path.exists():
        raise FileNotFoundError(f"TXT not found: {path}")

    logger.info(f"Reading TXT: {path}")

    try:

        with open(
            file_path,
            "r",
            encoding="utf-8"
        ) as file:

            text = clean_text(
                file.read()
            )

        return {

            "text": text,

            "preview": get_text_preview(text),

            "page_count": 1,

            "pages": [],

            "metadata": {}

        }

    except Exception as e:

        logger.exception("TXT Extraction Failed")

        raise RuntimeError(
            f"Failed to read TXT: {e}"
        )


# ======================================================
# Universal Document Reader
# ======================================================

def extract_document(path: str) -> dict:
    """
    Automatically detect document type
    and extract text.

    Supported:

    - PDF
    - DOCX
    - TXT
    """

    extension = Path(path).suffix.lower()

    if extension == ".pdf":
        return extract_text_from_pdf(path)

    if extension == ".docx":
        return extract_text_from_docx(path)

    if extension == ".txt":
        return extract_text_from_txt(path)

    raise ValueError(
        f"Unsupported file type: {extension}"
    )